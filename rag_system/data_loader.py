"""
Data Loader Module
Handles loading, preprocessing, and chunking of documents (PDF, TXT, DOCX)
"""

import os
from pathlib import Path
from typing import List, Dict, Any
import logging

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from pypdf import PdfReader
import docx

from config import config

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Loads documents from various formats and prepares them for RAG
    Supports: PDF, TXT, DOCX
    """
    
    def __init__(self):
        """Initialize document loader with text splitter"""
        # Calculate chunk size in characters (approximate: 1 word = 5 chars)
        chunk_size_chars = config.CHUNK_SIZE * 5
        chunk_overlap_chars = config.CHUNK_OVERLAP * 5
        
        # RecursiveCharacterTextSplitter is best for preserving semantic meaning
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size_chars,
            chunk_overlap=chunk_overlap_chars,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]  # Try to split at natural boundaries
        )
        
    def load_txt_file(self, file_path: str) -> str:
        """
        Load plain text file
        
        Args:
            file_path: Path to .txt file
            
        Returns:
            Text content as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            logger.info(f"✅ Loaded TXT file: {file_path} ({len(content)} chars)")
            return content
        except Exception as e:
            logger.error(f"❌ Error loading TXT file {file_path}: {e}")
            return ""
    
    def load_pdf_file(self, file_path: str) -> str:
        """
        Load PDF file and extract text
        
        Args:
            file_path: Path to .pdf file
            
        Returns:
            Extracted text content
        """
        try:
            reader = PdfReader(file_path)
            text = ""
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                text += page_text + "\n"
            
            logger.info(f"✅ Loaded PDF file: {file_path} ({len(reader.pages)} pages, {len(text)} chars)")
            return text
        except Exception as e:
            logger.error(f"❌ Error loading PDF file {file_path}: {e}")
            return ""
    
    def load_docx_file(self, file_path: str) -> str:
        """
        Load DOCX file and extract text
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Extracted text content
        """
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            logger.info(f"✅ Loaded DOCX file: {file_path} ({len(text)} chars)")
            return text
        except Exception as e:
            logger.error(f"❌ Error loading DOCX file {file_path}: {e}")
            return ""
    
    def load_document(self, file_path: str, metadata: Dict[str, Any] = None) -> str:
        """
        Load document based on file extension
        
        Args:
            file_path: Path to document
            metadata: Optional metadata to attach to document
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"❌ File not found: {file_path}")
            return ""
        
        # Route to appropriate loader based on extension
        extension = file_path.suffix.lower()
        
        if extension == '.txt':
            return self.load_txt_file(str(file_path))
        elif extension == '.pdf':
            return self.load_pdf_file(str(file_path))
        elif extension == '.docx':
            return self.load_docx_file(str(file_path))
        else:
            logger.error(f"❌ Unsupported file type: {extension}")
            return ""
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Split text into chunks suitable for RAG
        
        Args:
            text: Text to chunk
            metadata: Metadata to attach to each chunk
            
        Returns:
            List of LangChain Document objects with chunks
        """
        if not text.strip():
            return []
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create Document objects with metadata
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.copy() if metadata else {}
            doc_metadata['chunk_id'] = i
            doc_metadata['chunk_size'] = len(chunk)
            
            documents.append(Document(
                page_content=chunk,
                metadata=doc_metadata
            ))
        
        logger.info(f"✂️ Split text into {len(documents)} chunks (avg {sum(len(d.page_content) for d in documents) // len(documents)} chars/chunk)")
        return documents
    
    def load_and_chunk_document(self, file_path: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Complete pipeline: Load document → Chunk → Return Documents
        
        Args:
            file_path: Path to document
            metadata: Optional metadata
            
        Returns:
            List of chunked Document objects ready for embedding
        """
        # Add filename to metadata
        if metadata is None:
            metadata = {}
        metadata['source'] = str(Path(file_path).name)
        
        # Load document
        text = self.load_document(file_path, metadata)
        
        if not text:
            return []
        
        # Chunk text
        documents = self.chunk_text(text, metadata)
        
        return documents
    
    def load_directory(self, directory_path: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Load all supported documents from a directory
        
        Args:
            directory_path: Path to directory containing documents
            metadata: Optional metadata to apply to all documents
            
        Returns:
            List of all chunked documents from directory
        """
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            logger.error(f"❌ Directory not found: {directory}")
            return []
        
        all_documents = []
        supported_extensions = ['.txt', '.pdf', '.docx']
        
        # Find all supported files (recursively search subdirectories)
        files = [f for f in directory.rglob('*') if f.is_file() and f.suffix.lower() in supported_extensions]
        
        logger.info(f"📂 Found {len(files)} documents in {directory} (including subdirectories)")
        
        # Process each file
        for file_path in files:
            docs = self.load_and_chunk_document(str(file_path), metadata)
            all_documents.extend(docs)
        
        logger.info(f"✅ Loaded total {len(all_documents)} chunks from {len(files)} documents")
        return all_documents


# Convenience function
def load_documents(source: str, metadata: Dict[str, Any] = None) -> List[Document]:
    """
    Convenience function to load documents from file or directory
    
    Args:
        source: Path to file or directory
        metadata: Optional metadata
        
    Returns:
        List of Document objects
    """
    loader = DocumentLoader()
    
    source_path = Path(source)
    if source_path.is_dir():
        return loader.load_directory(str(source), metadata)
    else:
        return loader.load_and_chunk_document(str(source), metadata)


if __name__ == "__main__":
    """Test document loading"""
    print("🧪 Testing Document Loader\n")
    
    # Test with sample data
    test_file = config.DOCUMENTS_DIR / "sample_financial_data.txt"
    
    if not test_file.exists():
        # Create sample file for testing
        sample_content = """
        EMI (Equated Monthly Installment)
        
        EMI is a fixed payment amount made by a borrower to a lender at a specified date each calendar month. 
        EMIs are used to pay off both interest and principal each month, so that over a specified number of years, 
        the loan is fully paid off.
        
        The EMI formula is: EMI = [P x R x (1+R)^N]/[(1+R)^N-1]
        Where P = Principal loan amount, R = Rate of interest, N = Number of monthly installments
        
        UPI (Unified Payments Interface)
        
        UPI is an instant real-time payment system developed by National Payments Corporation of India (NPCI). 
        It facilitates inter-bank peer-to-peer and person-to-merchant transactions. 
        UPI is built over the IMPS infrastructure and allows you to instantly transfer money between any two parties' bank accounts.
        
        Fixed Deposit (FD)
        
        A Fixed Deposit is a financial instrument provided by banks which provides investors with a higher rate of interest 
        than a regular savings account, until the given maturity date. The tenure can range from 7 days to 10 years.
        """
        
        test_file.write_text(sample_content, encoding='utf-8')
        print(f"✅ Created sample file: {test_file}\n")
    
    # Load and chunk
    loader = DocumentLoader()
    documents = loader.load_and_chunk_document(
        str(test_file),
        metadata={"category": "financial_basics", "language": "en"}
    )
    
    print(f"\n📊 Results:")
    print(f"Total chunks: {len(documents)}")
    print(f"\nFirst chunk preview:")
    print(f"Content: {documents[0].page_content[:200]}...")
    print(f"Metadata: {documents[0].metadata}")
